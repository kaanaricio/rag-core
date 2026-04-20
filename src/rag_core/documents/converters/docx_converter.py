"""DOCX converter with hybrid text extraction + OCR fallback.

Uses python-docx for text extraction with heading style detection,
table extraction, and quality scoring.
"""

from __future__ import annotations

import asyncio
import io
import logging
from typing import Any, Dict, List, Optional

from .base import (
    ConversionResult,
    HybridConverter,
    render_markdown_table,
    score_text_quality,
)

logger = logging.getLogger(__name__)

# Heading style to markdown prefix mapping
_HEADING_MAP = (
    ("heading 1", "# "),
    ("heading 2", "## "),
    ("heading 3", "### "),
    ("heading", "#### "),
)


def _format_paragraph(para: Any) -> Optional[str]:
    """Convert a DOCX paragraph to a markdown line using style detection."""
    text = para.text.strip()
    if not text:
        return None

    style_name = (para.style.name or "").lower() if para.style else ""

    for keyword, prefix in _HEADING_MAP:
        if keyword in style_name:
            return "%s%s" % (prefix, text)

    if "list" in style_name:
        return "- %s" % text

    return text


def _format_table(table: Any) -> str:
    """Convert a DOCX table to markdown."""
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return render_markdown_table(rows)


def _extract_docx_figure_items(doc: Any) -> List[Dict[str, Any]]:
    """Extract lightweight figure metadata from DOCX inline shapes."""
    figures: List[Dict[str, Any]] = []
    for idx, shape in enumerate(getattr(doc, "inline_shapes", [])):
        figure_id = "fig:docx:%d" % (idx + 1)
        label = "DOCX Figure %d" % (idx + 1)
        description = "Embedded image extracted from DOCX."
        try:
            doc_pr = shape._inline.docPr  # type: ignore[attr-defined]
            alt_text = (
                getattr(doc_pr, "descr", None)
                or getattr(doc_pr, "title", None)
                or ""
            )
            if alt_text:
                description = str(alt_text).strip()
        except Exception:
            # Inline alt text is optional; keep the figure when descriptor lookup fails.
            pass

        figures.append(
            {
                "figure_id": figure_id,
                "page_index": 0,
                "label": label,
                "description": description,
                "metadata": {
                    "source": "docx:inline_shape",
                    "paragraph_index": idx + 1,
                },
            }
        )
    return figures


class DocxConverter(HybridConverter):
    """Converts DOCX files to markdown with heading style detection."""

    format_name = "docx"

    async def _try_extract(
        self,
        file_bytes: bytes,
        filename: str,
        mime_type: str,
    ) -> ConversionResult:
        """Extract text from DOCX using python-docx."""
        from docx import Document  # type: ignore[import-not-found]

        def _extract() -> ConversionResult:
            try:
                doc = Document(io.BytesIO(file_bytes))
            except Exception as exc:
                logger.warning("Failed to open DOCX %s: %s", filename, exc)
                return ConversionResult(
                    needs_ocr=True,
                    metadata={"parser": "local:python-docx", "error": str(exc)},
                )

            parts: List[str] = []

            for para in doc.paragraphs:
                line = _format_paragraph(para)
                if line:
                    parts.append(line)

            for table in doc.tables:
                md_table = _format_table(table)
                if md_table:
                    parts.append(md_table)

            figure_items = _extract_docx_figure_items(doc)
            if figure_items:
                parts.append("## Figures")
                for item in figure_items:
                    parts.append(
                        "- **%s**: %s"
                        % (item.get("label", "Figure"), item.get("description", ""))
                    )

            content = "\n\n".join(parts)
            quality = score_text_quality(content)

            metadata: Dict[str, Any] = {
                "parser": "local:python-docx",
                "needs_ocr": quality.char_count < 50,
            }
            if figure_items:
                metadata["figure_items"] = figure_items
                metadata["figure_count"] = len(figure_items)

            return ConversionResult(
                content=content,
                metadata=metadata,
                quality=quality,
            )

        return await asyncio.to_thread(_extract)
